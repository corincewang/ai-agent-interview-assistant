import asyncio
import re
import shutil
import subprocess
from tempfile import TemporaryDirectory
from pathlib import Path
from typing import Any
from uuid import UUID, uuid4

from app.domain.models import DocumentType, ParsedDocument, SourceSpan


DOCUMENT_NORMALIZATION_VERSION = "2026-05-12.unstructured-progressive-v1"
UNSTRUCTURED_SUFFIXES = {".pdf", ".doc", ".docx"}
MIN_EXTRACTABLE_PDF_CHARS = 300
PDF_PARSE_STRATEGIES = ("fast", "hi_res", "ocr_only")
PDF_PAGE_WINDOW_SIZE = 10
PDF_PAGE_WINDOW_THRESHOLD = 30


class LocalDocumentParsingTool:
    async def parse_document(
        self,
        file_path: Path,
        document_type: DocumentType,
        document_id: UUID | None = None,
    ) -> ParsedDocument:
        return await asyncio.to_thread(
            self._parse_document_sync,
            file_path,
            document_type,
            document_id,
        )

    async def parse_document_windows(
        self,
        file_path: Path,
        document_type: DocumentType,
        document_id: UUID | None = None,
    ):
        suffix = file_path.suffix.lower()
        resolved_document_id = document_id or uuid4()

        if suffix != ".pdf":
            yield await self.parse_document(
                file_path=file_path,
                document_type=document_type,
                document_id=resolved_document_id,
            )
            return

        page_count = _get_pdf_page_count(file_path)
        if page_count is None or page_count <= PDF_PAGE_WINDOW_THRESHOLD:
            yield await self.parse_document(
                file_path=file_path,
                document_type=document_type,
                document_id=resolved_document_id,
            )
            return

        for start_page, end_page in _iter_pdf_page_ranges(page_count):
            yield await asyncio.to_thread(
                self._parse_pdf_page_window,
                file_path,
                document_type,
                resolved_document_id,
                start_page,
                end_page,
            )

    def _parse_document_sync(
        self,
        file_path: Path,
        document_type: DocumentType,
        document_id: UUID | None = None,
    ) -> ParsedDocument:
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self._parse_pdf(file_path, document_type, document_id)

        if suffix in {".doc", ".docx"}:
            try:
                return self._parse_unstructured_document(file_path, document_type, document_id)
            except RuntimeError:
                if suffix == ".docx":
                    return self._parse_docx_fallback(
                        file_path=file_path,
                        document_type=document_type,
                        document_id=document_id,
                    )
                raise

        if suffix in {".txt", ".md"}:
            return self._parse_text(file_path, document_type, document_id)

        raise ValueError(f"Unsupported document type: {suffix}")

    def _parse_pdf(
        self,
        file_path: Path,
        document_type: DocumentType,
        document_id: UUID | None,
    ) -> ParsedDocument:
        page_count = _get_pdf_page_count(file_path)
        if page_count is not None and page_count > PDF_PAGE_WINDOW_THRESHOLD:
            return self._parse_pdf_by_page_windows(
                file_path=file_path,
                document_type=document_type,
                document_id=document_id,
                page_count=page_count,
            )

        return self._parse_pdf_with_progressive_unstructured(
            file_path=file_path,
            document_type=document_type,
            document_id=document_id,
        )

    def _parse_pdf_by_page_windows(
        self,
        file_path: Path,
        document_type: DocumentType,
        document_id: UUID | None,
        page_count: int,
    ) -> ParsedDocument:
        resolved_document_id = document_id or uuid4()
        parsed_windows: list[ParsedDocument] = []
        errors: list[str] = []

        with TemporaryDirectory(prefix="pdf-page-windows-") as output_dir:
            output_root = Path(output_dir)
            for start_page in range(1, page_count + 1, PDF_PAGE_WINDOW_SIZE):
                end_page = min(start_page + PDF_PAGE_WINDOW_SIZE - 1, page_count)
                try:
                    window_path = _create_pdf_page_window(
                        source_path=file_path,
                        output_root=output_root,
                        start_page=start_page,
                        end_page=end_page,
                    )
                    parsed_windows.append(
                        self._parse_unstructured_document(
                            file_path=window_path,
                            document_type=document_type,
                            document_id=resolved_document_id,
                            strategy="fast",
                            source_file_name=file_path.name,
                            page_number_offset=start_page - 1,
                            page_window={
                                "start_page": start_page,
                                "end_page": end_page,
                            },
                        )
                    )
                except Exception as exc:
                    errors.append(f"pages_{start_page}_{end_page}:{type(exc).__name__}")

        parsed_windows = [
            document
            for document in parsed_windows
            if _has_extractable_pdf_text(document)
        ]
        if not parsed_windows:
            raise RuntimeError(
                "unstructured could not parse usable text from PDF page windows. "
                f"attempts={errors}"
            )

        return _merge_parsed_pdf_windows(
            document_id=resolved_document_id,
            document_type=document_type,
            file_path=file_path,
            page_count=page_count,
            parsed_windows=parsed_windows,
            errors=errors,
        )

    def _parse_pdf_page_window(
        self,
        file_path: Path,
        document_type: DocumentType,
        document_id: UUID,
        start_page: int,
        end_page: int,
    ) -> ParsedDocument:
        with TemporaryDirectory(prefix="pdf-page-window-") as output_dir:
            window_path = _create_pdf_page_window(
                source_path=file_path,
                output_root=Path(output_dir),
                start_page=start_page,
                end_page=end_page,
            )
            return self._parse_unstructured_document(
                file_path=window_path,
                document_type=document_type,
                document_id=document_id,
                strategy="fast",
                source_file_name=file_path.name,
                page_number_offset=start_page - 1,
                page_window={
                    "start_page": start_page,
                    "end_page": end_page,
                },
            )

    def _parse_pdf_with_progressive_unstructured(
        self,
        file_path: Path,
        document_type: DocumentType,
        document_id: UUID | None,
    ) -> ParsedDocument:
        resolved_document_id = document_id or uuid4()
        errors: list[str] = []

        for strategy in PDF_PARSE_STRATEGIES:
            try:
                document = self._parse_unstructured_document(
                    file_path=file_path,
                    document_type=document_type,
                    document_id=resolved_document_id,
                    strategy=strategy,
                )
            except Exception as exc:
                errors.append(f"{strategy}:{type(exc).__name__}")
                continue

            if _has_extractable_pdf_text(document):
                document.metadata["parser_attempts"] = [*errors, f"{strategy}:success"]
                return document

            errors.append(f"{strategy}:low_text")

        raise RuntimeError(
            "unstructured could not parse usable text from PDF. "
            f"attempts={errors}"
        )

    def _parse_unstructured_document(
        self,
        file_path: Path,
        document_type: DocumentType,
        document_id: UUID | None,
        strategy: str | None = None,
        source_file_name: str | None = None,
        page_number_offset: int = 0,
        page_window: dict[str, int] | None = None,
    ) -> ParsedDocument:
        document_id = document_id or uuid4()
        suffix = file_path.suffix.lower()
        elements = _partition_document(file_path, strategy=strategy)
        normalized_elements = _normalize_unstructured_elements(
            elements,
            suffix,
            page_number_offset=page_number_offset,
        )
        raw_text = "\n\n".join(item["text"] for item in normalized_elements)
        source_spans = _build_source_spans(document_id, normalized_elements)
        sections = _extract_unstructured_sections(raw_text, normalized_elements)

        return ParsedDocument(
            id=document_id,
            document_type=document_type,
            raw_text=raw_text,
            source_spans=source_spans,
            metadata={
                "file_name": source_file_name or file_path.name,
                "file_suffix": suffix,
                "parser": "unstructured.partition.auto.partition",
                "parser_strategy": strategy or "auto",
                "normalization_version": DOCUMENT_NORMALIZATION_VERSION,
                "element_count": len(normalized_elements),
                "section_count": len(sections),
                "sections": sections,
                "element_types": _count_element_types(normalized_elements),
                "page_count": _count_pages(normalized_elements),
                "page_window": page_window,
            },
        )

    def _parse_docx_fallback(
        self,
        file_path: Path,
        document_type: DocumentType,
        document_id: UUID | None,
    ) -> ParsedDocument:
        document_id = document_id or uuid4()
        text = _extract_docx_text(file_path)
        sections = _extract_sections(text, source_format="word")

        return ParsedDocument(
            id=document_id,
            document_type=document_type,
            raw_text=text,
            source_spans=[
                SourceSpan(
                    document_id=document_id,
                    page_number=None,
                    start_char=0,
                    end_char=len(text),
                    text_excerpt=text[:500],
                )
            ],
            metadata={
                "file_name": file_path.name,
                "file_suffix": file_path.suffix.lower(),
                "parser": "unstructured.partition.auto.partition",
                "parser_strategy": "docx-python-docx-fallback",
                "normalization_version": DOCUMENT_NORMALIZATION_VERSION,
                "section_count": len(sections),
                "sections": sections,
            },
        )

    def _parse_text(
        self,
        file_path: Path,
        document_type: DocumentType,
        document_id: UUID | None,
    ) -> ParsedDocument:
        document_id = document_id or uuid4()
        suffix = file_path.suffix.lower()
        text = _clean_structured_text(
            file_path.read_text(encoding="utf-8"),
            source_format="markdown" if suffix == ".md" else "text",
        )
        sections = _extract_sections(
            text,
            source_format="markdown" if suffix == ".md" else "text",
        )

        return ParsedDocument(
            id=document_id,
            document_type=document_type,
            raw_text=text,
            source_spans=[
                SourceSpan(
                    document_id=document_id,
                    page_number=None,
                    start_char=0,
                    end_char=len(text),
                    text_excerpt=text[:500],
                )
            ],
            metadata={
                "file_name": file_path.name,
                "file_suffix": suffix,
                "normalization_version": DOCUMENT_NORMALIZATION_VERSION,
                "section_count": len(sections),
                "sections": sections,
            },
        )


def _partition_document(file_path: Path, strategy: str | None = None) -> list[Any]:
    from unstructured.partition.auto import partition

    kwargs = {"filename": str(file_path)}
    if strategy is not None:
        kwargs["strategy"] = strategy
    return list(partition(**kwargs))


def _extract_docx_text(file_path: Path) -> str:
    from docx import Document

    document = Document(str(file_path))
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    return _clean_structured_text("\n\n".join(paragraphs), source_format="word")


def _has_extractable_pdf_text(document: ParsedDocument) -> bool:
    return len(document.raw_text.strip()) >= MIN_EXTRACTABLE_PDF_CHARS


def _get_pdf_page_count(file_path: Path) -> int | None:
    executable = shutil.which("pdfinfo")
    if executable is None:
        return None

    completed = subprocess.run(
        [executable, str(file_path)],
        check=False,
        capture_output=True,
        text=True,
    )
    if completed.returncode != 0:
        return None

    for line in completed.stdout.splitlines():
        if line.startswith("Pages:"):
            return int(line.split(":", 1)[1].strip())

    return None


def _create_pdf_page_window(
    source_path: Path,
    output_root: Path,
    start_page: int,
    end_page: int,
) -> Path:
    pdfseparate = shutil.which("pdfseparate")
    pdfunite = shutil.which("pdfunite")
    if pdfseparate is None or pdfunite is None:
        raise RuntimeError("Poppler pdfseparate/pdfunite are required for page-window parsing.")

    window_dir = output_root / f"pages_{start_page}_{end_page}"
    window_dir.mkdir(parents=True, exist_ok=True)
    page_pattern = window_dir / "page-%d.pdf"
    window_path = output_root / f"window_{start_page}_{end_page}.pdf"

    separate_result = subprocess.run(
        [
            pdfseparate,
            "-f",
            str(start_page),
            "-l",
            str(end_page),
            str(source_path),
            str(page_pattern),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    if separate_result.returncode != 0:
        raise RuntimeError(separate_result.stderr.strip() or separate_result.stdout.strip())

    page_files = [window_dir / f"page-{page_number}.pdf" for page_number in range(start_page, end_page + 1)]
    unite_result = subprocess.run(
        [pdfunite, *[str(page_file) for page_file in page_files], str(window_path)],
        check=False,
        capture_output=True,
        text=True,
    )
    if unite_result.returncode != 0:
        raise RuntimeError(unite_result.stderr.strip() or unite_result.stdout.strip())

    return window_path


def _iter_pdf_page_ranges(page_count: int) -> list[tuple[int, int]]:
    return [
        (
            start_page,
            min(start_page + PDF_PAGE_WINDOW_SIZE - 1, page_count),
        )
        for start_page in range(1, page_count + 1, PDF_PAGE_WINDOW_SIZE)
    ]


def _merge_parsed_pdf_windows(
    document_id: UUID,
    document_type: DocumentType,
    file_path: Path,
    page_count: int,
    parsed_windows: list[ParsedDocument],
    errors: list[str],
) -> ParsedDocument:
    raw_parts: list[str] = []
    source_spans: list[SourceSpan] = []
    sections: list[dict[str, object]] = []
    element_types: dict[str, int] = {}
    page_windows: list[dict[str, int]] = []
    cursor = 0
    total_elements = 0

    for window in parsed_windows:
        if raw_parts:
            cursor += 2

        raw_parts.append(window.raw_text)
        source_spans.extend(_shift_source_spans(window.source_spans, cursor))
        sections.extend(_shift_sections(window.metadata.get("sections", []), cursor))

        total_elements += int(window.metadata.get("element_count") or 0)
        page_window = window.metadata.get("page_window")
        if isinstance(page_window, dict):
            page_windows.append(page_window)

        for key, count in dict(window.metadata.get("element_types") or {}).items():
            element_types[str(key)] = element_types.get(str(key), 0) + int(count)

        cursor += len(window.raw_text)

    raw_text = "\n\n".join(raw_parts)

    return ParsedDocument(
        id=document_id,
        document_type=document_type,
        raw_text=raw_text,
        source_spans=source_spans,
        metadata={
            "file_name": file_path.name,
            "file_suffix": file_path.suffix.lower(),
            "parser": "unstructured.partition.auto.partition",
            "parser_strategy": "fast-page-windowed",
            "normalization_version": DOCUMENT_NORMALIZATION_VERSION,
            "element_count": total_elements,
            "section_count": len(sections),
            "sections": sections,
            "element_types": element_types,
            "page_count": page_count,
            "page_window_size": PDF_PAGE_WINDOW_SIZE,
            "page_windows": page_windows,
            "parser_attempts": [*errors, "page_windows:success"],
        },
    )


def _shift_source_spans(source_spans: list[SourceSpan], offset: int) -> list[SourceSpan]:
    return [
        SourceSpan(
            document_id=source_span.document_id,
            page_number=source_span.page_number,
            start_char=_shift_optional_int(source_span.start_char, offset),
            end_char=_shift_optional_int(source_span.end_char, offset),
            text_excerpt=source_span.text_excerpt,
        )
        for source_span in source_spans
    ]


def _shift_sections(sections: object, offset: int) -> list[dict[str, object]]:
    shifted_sections: list[dict[str, object]] = []
    if not isinstance(sections, list):
        return shifted_sections

    for section in sections:
        if not isinstance(section, dict):
            continue

        shifted_section = dict(section)
        if isinstance(shifted_section.get("start_char"), int):
            shifted_section["start_char"] = int(shifted_section["start_char"]) + offset
        if isinstance(shifted_section.get("end_char"), int):
            shifted_section["end_char"] = int(shifted_section["end_char"]) + offset
        shifted_sections.append(shifted_section)

    return shifted_sections


def _shift_optional_int(value: int | None, offset: int) -> int | None:
    if value is None:
        return None
    return value + offset


def _normalize_unstructured_elements(
    elements: list[Any],
    suffix: str,
    page_number_offset: int = 0,
) -> list[dict[str, object]]:
    normalized: list[dict[str, object]] = []
    cursor = 0

    for index, element in enumerate(elements):
        text = _clean_unstructured_text(str(element), suffix)
        if not text:
            continue

        category = _element_category(element)
        metadata = getattr(element, "metadata", None)
        page_number = getattr(metadata, "page_number", None) if metadata is not None else None
        if page_number is not None:
            page_number = int(page_number) + page_number_offset
        start = cursor
        end = start + len(text)

        normalized.append(
            {
                "index": index,
                "text": text,
                "category": category,
                "page_number": page_number,
                "start_char": start,
                "end_char": end,
            }
        )
        cursor = end + 2

    return normalized


def _clean_unstructured_text(text: str, suffix: str) -> str:
    if suffix == ".pdf":
        return _clean_pdf_like_text(text)
    return _clean_structured_text(text, source_format="word")


def _clean_pdf_like_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _merge_isolated_symbol_lines(text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[，。！？；：（）、])", "", text)
    text = re.sub(r"(?<=[，。！？；：（）、])\s+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"\s+([，。！？；：、,.!?;:%）】》」』])", r"\1", text)
    text = re.sub(r"([（【《「『])\s+", r"\1", text)
    text = re.sub(r"\s*([+/#&])\s*", r"\1", text)
    text = re.sub(r"(?<=[A-Za-z0-9])\s*\n+\s*(?=[A-Za-z0-9])", " ", text)
    text = re.sub(r"(?<=[A-Za-z0-9])\s*\n+\s*(?=[\u4e00-\u9fff])", " ", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s*\n+\s*(?=[A-Za-z0-9])", " ", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _merge_isolated_symbol_lines(text: str) -> str:
    lines = text.splitlines()
    merged: list[str] = []
    symbol_pattern = re.compile(r"^[“”‘’\"'：:，,。.!！?？、；;（）()\[\]【】《》+\-/|]+$")

    for raw_line in lines:
        line = raw_line.strip()
        if line and symbol_pattern.match(line) and merged:
            merged[-1] = f"{merged[-1]}{line}"
            continue

        if merged and merged[-1] and line:
            previous = merged[-1]
            if previous.endswith(("“", "‘", "\"", "'", "：", ":", "+", "/", "-", "|")):
                merged[-1] = f"{previous}{line}"
                continue

        merged.append(raw_line)

    return "\n".join(merged)


def _element_category(element: Any) -> str:
    category = getattr(element, "category", None)
    if category:
        return str(category)
    return type(element).__name__


def _build_source_spans(
    document_id: UUID,
    normalized_elements: list[dict[str, object]],
) -> list[SourceSpan]:
    return [
        SourceSpan(
            document_id=document_id,
            page_number=_optional_int(element.get("page_number")),
            start_char=int(element["start_char"]),
            end_char=int(element["end_char"]),
            text_excerpt=str(element["text"])[:500],
        )
        for element in normalized_elements
    ]


def _extract_unstructured_sections(
    raw_text: str,
    normalized_elements: list[dict[str, object]],
) -> list[dict[str, object]]:
    heading_elements = [
        element
        for element in normalized_elements
        if str(element["category"]).lower() == "title"
        or _parse_heading(str(element["text"]).strip(), source_format="word") is not None
    ]
    sections: list[dict[str, object]] = []

    for index, heading in enumerate(heading_elements):
        next_start = (
            int(heading_elements[index + 1]["start_char"])
            if index + 1 < len(heading_elements)
            else len(raw_text)
        )
        sections.append(
            {
                "title": str(heading["text"])[:120],
                "level": 2 if str(heading["category"]).lower() == "title" else 3,
                "start_char": int(heading["start_char"]),
                "end_char": next_start,
                "source": "unstructured",
                "page_number": heading.get("page_number"),
            }
        )

    return sections


def _count_element_types(normalized_elements: list[dict[str, object]]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for element in normalized_elements:
        category = str(element["category"])
        counts[category] = counts.get(category, 0) + 1
    return counts


def _count_pages(normalized_elements: list[dict[str, object]]) -> int | None:
    pages = {
        page_number
        for page_number in (element.get("page_number") for element in normalized_elements)
        if page_number is not None
    }
    return len(pages) if pages else None


def _optional_int(value: object) -> int | None:
    if value is None:
        return None
    return int(value)


def _clean_structured_text(text: str, source_format: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(?<=[\u4e00-\u9fff])[ \t]+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])[ \t]+(?=[，。！？；：（）、])", "", text)
    text = re.sub(r"(?<=[，。！？；：（）、])[ \t]+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)

    if source_format == "markdown":
        text = _normalize_markdown_headings(text)

    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _normalize_markdown_headings(text: str) -> str:
    lines = text.splitlines()
    normalized: list[str] = []

    for line in lines:
        heading_match = re.match(r"^(#{1,6})([^#\s].*)$", line)
        if heading_match:
            normalized.append(f"{heading_match.group(1)} {heading_match.group(2).strip()}")
        else:
            normalized.append(line)

    return "\n".join(normalized)


def _extract_sections(text: str, source_format: str) -> list[dict[str, object]]:
    heading_spans = _find_heading_spans(text, source_format)
    sections: list[dict[str, object]] = []

    for index, heading in enumerate(heading_spans):
        next_start = (
            heading_spans[index + 1]["start_char"]
            if index + 1 < len(heading_spans)
            else len(text)
        )
        sections.append(
            {
                "title": heading["title"],
                "level": heading["level"],
                "start_char": heading["start_char"],
                "end_char": next_start,
            }
        )

    return sections


def _find_heading_spans(text: str, source_format: str) -> list[dict[str, object]]:
    headings: list[dict[str, object]] = []
    cursor = 0

    for line in text.splitlines(keepends=True):
        stripped = line.strip()
        if not stripped:
            cursor += len(line)
            continue

        heading = _parse_heading(stripped, source_format)
        if heading is not None:
            headings.append(
                {
                    "title": heading["title"],
                    "level": heading["level"],
                    "start_char": cursor,
                }
            )

        cursor += len(line)

    return headings


def _parse_heading(line: str, source_format: str) -> dict[str, object] | None:
    if source_format == "markdown":
        markdown_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if markdown_match:
            return {
                "title": markdown_match.group(2).strip(),
                "level": len(markdown_match.group(1)),
            }

    numbered_match = re.match(
        r"^((第[一二三四五六七八九十\d]+[章节部分])|([一二三四五六七八九十\d]+[、.．]))\s*(.{1,80})$",
        line,
    )
    if numbered_match:
        if re.search(r"[。！？.!?；;]$", line):
            return None
        return {"title": line, "level": 2}

    if source_format != "pdf" and _looks_like_short_heading(line):
        return {"title": line, "level": 3}

    return None


def _looks_like_short_heading(line: str) -> bool:
    if len(line) > 48:
        return False
    if re.search(r"[。！？.!?；;，,]$", line):
        return False
    if re.match(r"^[-*+]\s+", line):
        return False
    if re.match(r"^\d+[\).]\s+", line):
        return False

    heading_keywords = (
        "项目",
        "经历",
        "技能",
        "教育",
        "面试",
        "问题",
        "答案",
        "知识",
        "React",
        "JavaScript",
        "TypeScript",
        "Python",
        "Agent",
        "RAG",
        "Backend",
        "Frontend",
    )
    return any(keyword in line for keyword in heading_keywords)
